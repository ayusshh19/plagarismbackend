from django.shortcuts import render
import subprocess
import os.path
import sys
from shutil import copyfile
import pdfx
from PyPDF2 import PdfMerger
import io
from PIL import Image  
import tempfile 
import fitz
import pdfplumber
import docx
import PyPDF2
import zipfile
from django.conf import settings
from io import BytesIO
import pytesseract  #install this on cpu before pip install,  also make sure folder for it is in (86)program file
from PIL import Image 
import pathlib
from pathlib import Path
import fitz  #trying this instead of poppler
import os
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
import json
from django.http import HttpResponse

# Create your views here.
def compress(input_file_path, output_file_path, power=0):
    """Function to compress PDF via Ghostscript command line interface"""
    quality = {
        0: '/default',
        1: '/prepress',
        2: '/printer',
        3: '/ebook',
        4: '/screen'
    }

    # Basic controls
    # Check if valid path
    if not os.path.isfile(input_file_path):
        print("Error: invalid path for input PDF file")
        sys.exit(1)

    # Check if file is a PDF by extension
    if input_file_path.split('.')[-1].lower() != 'pdf':
        print("Error: input file is not a PDF")
        sys.exit(1)

    print("Compress PDF...")
    initial_size = os.path.getsize(input_file_path)
    subprocess.call(['gs', '-sDEVICE=pdfwrite', '-dCompatibilityLevel=1.4',
                    '-dPDFSETTINGS={}'.format(quality[power]),
                    '-dNOPAUSE', '-dQUIET', '-dBATCH',
                    '-sOutputFile={}'.format(output_file_path),
                     input_file_path]
    )
    final_size = os.path.getsize(output_file_path)
    ratio = 1 - (final_size / initial_size)
    print("Compression by {0:.0%}.".format(ratio))
    print("Final file size is {0:.1f}MB".format(final_size / 1000000))
    print("Done.")
    
def serialize_set(obj):
    if isinstance(obj, set):
        return list(obj)
    raise TypeError('Object of type %s is not JSON serializable' % type(obj).__name__)

@api_view(['GET','POST'])
def getreference(request):
    if request.method=='POST':
        pdf = pdfx.PDFx(request.data['url'])
        metadata = pdf.get_metadata()
        references_list = pdf.get_references()
        references_dict = pdf.get_references_as_dict()
        references_dict=dict(references_dict)
        metadata=dict(metadata)
        print(references_dict,metadata,references_list)
        return Response({'msg':'my post request','referencedict':references_dict,'referencemetadata':metadata},status=status.HTTP_200_OK)
    return Response({'msg':'get request'},status=status.HTTP_200_OK)
    
@api_view(['GET','POST'])    
def mergepdf(request):
 if request.method=='POST':
             # Get the uploaded file from the request object
        uploaded_file1 = request.FILES.getlist('files')[0]
        uploaded_file2 = request.FILES.getlist('files')[1]
        # Load the file into a PyPDF2.PdfFileReader object
        pdf_reader1 = PyPDF2.PdfReader(BytesIO(uploaded_file1.read()))
        pdf_reader2 = PyPDF2.PdfReader(BytesIO(uploaded_file2.read()))
        pdfs = [pdf_reader1, pdf_reader2]

        merger = PdfMerger()

        for pdf in pdfs:
            merger.append(pdf)
        output = io.BytesIO()

    # Write the merged PDF to the output file
        merger.write(output)
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="example.pdf"'
        response.write(output.getvalue())
        return response
        # return Response({'msg':'my post request','pdf':merger},status=status.HTTP_200_OK)
 return Response({'msg':'get request'},status=status.HTTP_200_OK)
BASE_DIR = Path(__file__).resolve().parent.parent


@api_view(['GET','POST'])  
def extractimage(request):
    if request.method=='POST':
        uploaded_file1 = request.FILES['files']
        if not uploaded_file1 or uploaded_file1.size == 0:
            return HttpResponse('PDF file is empty', status=400)
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            # Save uploaded file to temporary file on server
            for chunk in uploaded_file1.chunks():
                temp_file.write(chunk)
        try:
            # Open PDF file using fitz
            with fitz.open(temp_file.name) as doc:
                # iterate over PDF pages
                for page_index in range(len(doc)):

                    # get the page itself
                    page = doc[page_index]
                    image_list = page.get_images()

                    # printing number of images found in this page
                    if image_list:
                        print(f"[+] Found a total of {len(image_list)} images in page {page_index}")
                    else:
                        print("[!] No images found on page", page_index)
                    for image_index, img in enumerate(page.get_images(), start=1):
                        # get the XREF of the image
                        xref = img[0]

                        # extract the image bytes   
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]

                        # get the image extension
                        image_ext = base_image["ext"]

                        img = Image.open(io.BytesIO(image_bytes))
                        img.save(open(f"image{page_index + 1}_{image_index}.{image_ext}", "wb")) 
        except (OSError, ValueError) as e:
            print(f"Error opening PDF file: {e}")
        finally:
            # Delete temporary file
            os.unlink(temp_file.name)
        # Create a Django HTTP response containing the image
        zip_filename = 'images.zip'
        zip_path = os.path.join(settings.MEDIA_ROOT, zip_filename)
        with zipfile.ZipFile(zip_path, 'w') as zip_file:
            for dirpath, dirnames, filenames in os.walk(BASE_DIR):
                count=0
                for filename in filenames:
                    if filename.lower().endswith('.jpeg') or filename.lower().endswith('.jpg') or filename.lower().endswith('.png'):
                        file_path = os.path.join(dirpath, filename)
                        zip_file.write(file_path,arcname=os.path.relpath(file_path, BASE_DIR))
                        count+=1
                        if(count>12):
                            break
                break

        # Read the zip archive into memory and delete the file from disk
        with open(zip_path, 'rb') as zip_file:
            zip_content = zip_file.read()
        os.remove(zip_path)
    
        # Create the HTTP response
        response = HttpResponse(zip_content, content_type='application/zip')
        response['Content-Disposition'] = f'attachment; filename="{zip_filename}"'
        return response
        return Response({'msg':'my post request','pdf':merger},status=status.HTTP_200_OK)
    return Response({'msg':'get request'},status=status.HTTP_200_OK)
        

           

@api_view(['GET','POST'])
def pdftoword(request):
    if request.method=='POST':
        # Open PDF file and extract text with pdfplumber
        with pdfplumber.open('plagarisimdetectionofimages.pdf') as pdf:
            text = ''
            for page in pdf.pages:
                text += page.extract_text()

        # Create a Word document with docx
        doc = docx.Document()

        # Add text to the Word document with formatting and alignment
        for line in text.split('\n'):
            if line.strip():
                # Check if the line is centered
                if line.center(len(line)).strip() == line.strip():
                    p = doc.add_paragraph(line)
                # Check if the line is right-aligned
                elif line.rstrip() == line.strip():
                    p = doc.add_paragraph(line, style='Right')
                # Otherwise, add the line as left-aligned
                else:
                    p = doc.add_paragraph(line)

        # Save the Word document
        doc.save('example1.docx')
    
pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files (x86)\\Tesseract-OCR\\tesseract.exe' 
working_directory = Path(__file__).absolute().parent
scanned_pdf = working_directory / 'plagarisimdetectionofimages.pdf'
@api_view(['GET','POST'])
def scannedpdf(request):
    if request.method=='POST':
        uploaded_file1 = request.FILES['files']
        # Load the file into a PyPDF2.PdfFileReader object
        pdf_reader1 = PyPDF2.PdfReader(BytesIO(uploaded_file1.read()))
        if not uploaded_file1 or uploaded_file1.size == 0:
            return HttpResponse('PDF file is empty', status=400)
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            # Save uploaded file to temporary file on server
            for chunk in uploaded_file1.chunks():
                temp_file.write(chunk)
        try:
            # Open PDF file using fitz
            content=''
            with fitz.open(temp_file.name) as doc:
                for i in range(len(doc)):
                    page = doc.load_page(i)  # number of page
                    pix = page.get_pixmap()
                    output = f"outfile{i}.png"
                    pix.save(output)
                    text=pytesseract.image_to_string(output) 
                    content+=text
                    os.remove(f"output{i}.png")
        except (OSError, ValueError) as e:
            print(f"Error opening PDF file: {e}")
        finally:
            # Delete temporary file
            os.unlink(temp_file.name)
        
        return Response({'msg':'my post request','textdata':content},status=status.HTTP_200_OK)
    return Response({'msg':'get request'},status=status.HTTP_200_OK)