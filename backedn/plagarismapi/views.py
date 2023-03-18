from django.shortcuts import render
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from difflib import SequenceMatcher
from django.db.models.signals import post_save
from django.dispatch import receiver
from .models import Documents
import os
from PIL import Image
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from pathlib import Path
import pytesseract
import PyPDF2
import docx
from io import BytesIO
import cv2
import pytesseract
import matplotlib.image as mpimg
import matplotlib.pyplot as plt
import numpy as np
from .cursive.image_straighten import image_straighten
from .cursive.center_align import center_align
from .cursive.segmentation import segmentation
from .cursive.model_build import model_build
from .cursive.recognition import recognition
# Build paths inside the project like this: BASE_DIR / 'subdir'.
BASE_DIR = Path(__file__).resolve().parent.parent
# Create your views here.
pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files (x86)\\Tesseract-OCR\\tesseract.exe'
@api_view(['GET','POST'])
def home(request):
    return Response({'msg':'hello user this is our plagarism'},status=status.HTTP_200_OK)

def wordplagariser(file1,file2):
    with open(file1,'r', encoding='utf-8', errors='ignore') as file1 ,open(file2,'r', encoding='utf-8', errors='ignore') as file2:
        filldata=file1.read()
        fill2data=file2.read()
        similarity=SequenceMatcher(None,filldata,fill2data).ratio()
        return similarity*100
        
MEDIA_ROOT=BASE_DIR/'media'  
   
def get_media_files(topic,type):
    media_files = []
    checkdir=False
    index=0
    currentno=0
    for root, dirs, files in os.walk(MEDIA_ROOT):
        # print(f'dirs {dirs}  files {files}')
        if topic in dirs:
            checkdir=True
            index=dirs.index(topic)
            # print(f'index {index}')
            
        if checkdir :
            if index==currentno:
                # print(f'media file length {len(media_files)}')
                for file in files:
                    media_files.append(os.path.join(root, file))
                if len(media_files)>1:
                    checkdir=False
                    currentno+=1
               
        
    return media_files

@receiver(post_save, sender=Documents)
def do_something(sender, instance, **kwargs):
    scoreindex=[]
    if instance.docname:
        File_List=get_media_files(instance.topic,'worddocument')
        setfile=str(instance.docname).split('/')
        setfilenew="\\".join(setfile)
        print(setfilenew)
        filename=os.path.join(MEDIA_ROOT, str(setfilenew))
        for item in File_List:
            print(f'item {item} filename {filename} is equeal {item==filename}')
            if item == filename:
                scoreindex.append(0)
                continue
            else:
                result=wordplagariser(filename,item)
                scoreindex.append(result)
        try:
          if len(scoreindex)>0:
              print(Documents.objects.filter(id=instance.id))
              success=Documents.objects.filter(id=instance.id).update(plagarism=max(scoreindex))
              print(success)
              print(scoreindex)
              print(f'success {success}')
        except :
          print('An exception occurred')
          
@api_view(['GET','POST'])
def process_image(request):
    if request.method == 'POST':
        print(request.data)
        images = request.FILES.getlist('images')
        imagetextlist=[]
        for image in images:
            image = Image.open(image)
            text = pytesseract.image_to_string(image)
            imagetextlist.append(text)
        similarity=SequenceMatcher(None,imagetextlist[0],imagetextlist[1]).ratio()
        return Response({'msg':'my post request','Score':similarity*100},status=status.HTTP_200_OK)
    return Response({'msg':'get request'},status=status.HTTP_200_OK)

@api_view(['GET','POST'])
def wordtotext(request):
    if request.method == 'POST':
        print(request.data)
        document1 = docx.Document(request.FILES.getlist('worddocuments')[0])
        document2 = docx.Document(request.FILES.getlist('worddocuments')[1])
        firstdoc=''
        seconddoc=''
        for paragraph in document1.paragraphs:
            firstdoc+=paragraph.text
        for paragraph in document2.paragraphs:
            seconddoc+=paragraph.text
        similarity=SequenceMatcher(None,firstdoc,seconddoc).ratio()
        return Response({'msg':'my post request','Score':similarity*100},status=status.HTTP_200_OK)
    return Response({'msg':'get request'},status=status.HTTP_200_OK)

@api_view(['GET','POST'])
def extract_text_from_pdf(request):
    if request.method=='POST':
        # Get the uploaded file from the request object
        uploaded_file1 = request.FILES.getlist('document')[0]
        uploaded_file2 = request.FILES.getlist('document')[1]
        # Load the file into a PyPDF2.PdfFileReader object
        pdf_reader1 = PyPDF2.PdfReader(BytesIO(uploaded_file1.read()))
        pdf_reader2 = PyPDF2.PdfReader(BytesIO(uploaded_file2.read()))
        

        # Extract the text from each page of the PDF file
        text1 = ''
        text2 = ''
        for i in range(len(pdf_reader1.pages)):
            text1 += pdf_reader1.pages[i].extract_text()
        print(text1)
        for i in range(len(pdf_reader2.pages)):
            text2 += pdf_reader1.pages[i].extract_text()
        similarity=SequenceMatcher(None,text1,text2).ratio()
        return Response({'msg':'my post request','Score':similarity*100},status=status.HTTP_200_OK)
    return Response({'msg':'get request'},status=status.HTTP_200_OK)

@api_view(['GET','POST'])
def imagefeature(request):
    if request.method == 'POST':
        print(request.data)
        images = request.FILES.getlist('images')
        imagetextlist=[]
        for image in images:
            img_bytes = image.read()
            nparr = np.frombuffer(img_bytes, np.uint8)

            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            imagetextlist.append(img)
        template_gray = cv2.cvtColor(imagetextlist[0], cv2.COLOR_BGR2GRAY)
        target_gray = cv2.cvtColor(imagetextlist[1], cv2.COLOR_BGR2GRAY)

        sift = cv2.SIFT_create()
        kp1, des1 = sift.detectAndCompute(template_gray, None)
        kp2, des2 = sift.detectAndCompute(target_gray, None)

        bf = cv2.BFMatcher()
        matches = bf.match(des1, des2)
        matches = sorted(matches, key=lambda x: x.distance)

        src_pts = np.float32([kp1[m.queryIdx].pt for m in matches]).reshape(-1, 1, 2)
        dst_pts = np.float32([kp2[m.trainIdx].pt for m in matches]).reshape(-1, 1, 2)
        M, mask = cv2.findHomography(src_pts, dst_pts, cv2.RANSAC, 5.0)

        h, w = template_gray.shape
        warped_template = cv2.warpPerspective(template_gray, M, (w, h))
        
        if template_gray.dtype != target_gray.dtype:
            target_gray = target_gray.astype(template_gray.dtype)
            
        if template_gray.shape[0] > target_gray.shape[0] or template_gray.shape[1] > target_gray.shape[1]:
            template_gray = cv2.resize(template_gray, (target_gray.shape[1], target_gray.shape[0]))

        similarity_metric = cv2.matchTemplate(target_gray,template_gray, cv2.TM_CCOEFF_NORMED)
        similarity_score = similarity_metric.max()

        threshold = 0.8
        print(f'similarity {similarity_score}')
        if similarity_score >= threshold:
            print("Plagiarism detected!")
        else:
            print("No plagiarism detected.")
        return Response({'msg':'my post request','Score':similarity_score*100},status=status.HTTP_200_OK)
    return Response({'msg':'get request'},status=status.HTTP_200_OK)

build_model = 0 # Model already exists
@api_view(['GET','POST'])
def handwritten(request):
    if request.method == 'POST':
        print(request.data)
        images = request.FILES.getlist('images')
        imagetextlist=[]
        for image in images:
            img_bytes = image.read()
            nparr = np.frombuffer(img_bytes, np.uint8)

            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            imagetextlist.append(img)
        result=[]
        for img in imagetextlist:
            image_straighten(img)
            segmentation(img)
            folder='plagarismapi/cursive/result/characters/'
            center_align(folder)

#            Fix model building and recognition later
            if (build_model != 0):
                model_build()
            rec_char = []
            rec_char = recognition()
            result.append(rec_char)
        print(result)
        return Response({'msg':'my post request','Score':1*100},status=status.HTTP_200_OK)
    return Response({'msg':'get request'},status=status.HTTP_200_OK)

build_model = 0 # Model already exists
@api_view(['GET','POST'])
def normalhandwritten(request):
    if request.method == 'POST':
        print(request.data)
        images = request.FILES.getlist('images')
        imagetextlist=[]
        for image in images:
            image = Image.open(image)
            text = pytesseract.image_to_string(image)
            imagetextlist.append(text)
            print(text)
        similarity=SequenceMatcher(None,imagetextlist[0],imagetextlist[1]).ratio()
        return Response({'msg':'my post request','Score':similarity*100},status=status.HTTP_200_OK)
    return Response({'msg':'get request'},status=status.HTTP_200_OK)